from docx import Document


def replace_text_in_docx(file_path, old_text, new_text, output_path):
    # Open the document
    doc = Document(file_path)

    # Function to replace text in paragraphs, tables, headers, and footers
    def replace_in_paragraphs(paragraphs):
        for paragraph in paragraphs:
            if old_text in paragraph.text:
                inline = paragraph.runs
                for run in inline:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

    # 1. Replace text in regular paragraphs
    replace_in_paragraphs(doc.paragraphs)

    # 2. Replace text in tables (for each table and its cells)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

    # 3. Replace text in headers
    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs)

    # 4. Replace text in footers
    for section in doc.sections:
        replace_in_paragraphs(section.footer.paragraphs)

    # Save the modified document to a new file
    doc.save(output_path)


if __name__ == "__main__":
    # Step 3: Run the script to replace text and export to PDF
    docx_path = 'cv.docx'  # Replace with your file path
    old_text = 'Daan Zeeuwe is an ambitious researcher with a diverse educational background currently focusing on developing scalable products in startups. His focus has shifted from embedded programming to (robotic) software development and Deep Learning. His ambition includes working towards the next phase of human-like intelligence, using the literature from neuroscience and psychology as inspiration for deep learning and machine learning approaches for developing artificial general intelligence. '  # Replace with the text you want to find
    new_text = 'Daan Zeeuwe is a driven researcher with a broad educational background, currently concentrating on creating scalable products in startups. His expertise has evolved from embedded programming to (robotic) software development and deep learning. He is dedicated to advancing toward the next phase of human-like intelligence, drawing inspiration from neuroscience and psychology to inform deep learning and machine learning techniques in the pursuit of artificial general intelligence.'  # Replace with the text you want to replace it with
    output_path = 'modified_text.docx'

    replace_text_in_docx(docx_path, old_text, new_text, output_path)
